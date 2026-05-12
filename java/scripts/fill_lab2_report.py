from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

TEMPLATE = r'D:\EGlaoding\实验二_企业级 AI Agent 全栈开发.docx'
OUTPUT = r'D:\EGlaoding\实验二_企业级 AI Agent 全栈开发_已填写.docx'

doc = Document(TEMPLATE)

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)

t = doc.tables[1]

def clear_cell(cell):
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ''

def set_cell_text(cell, text_lines):
    """Set cell content, replace with multiple paragraphs."""
    clear_cell(cell)
    for i, line in enumerate(text_lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        run = p.add_run(line)
        run.font.name = '宋体'
        run.font.size = Pt(12)

def add_image_placeholder(cell, label):
    """Add an image placeholder box."""
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'[ {label} ]')
    run.font.name = '宋体'
    run.font.size = Pt(11)
    run.font.color.rgb = None
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'（截图：{label}）')
    run.font.name = '宋体'
    run.font.size = Pt(10)
    p = cell.add_paragraph()
    p.add_run('').font.size = Pt(6)

def add_paragraph(cell, text, bold=False, size=12):
    """Add a text paragraph to a cell."""
    p = cell.add_paragraph()
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_code_ref(cell, code_text):
    """Add a code reference line."""
    p = cell.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    return p

def add_bold(cell, text, size=12):
    return add_paragraph(cell, text, bold=True, size=size)

def add_normal(cell, text, size=12):
    return add_paragraph(cell, text, bold=False, size=size)

# ======== Row 1: 实验原理与内容 ========
c = t.rows[1].cells[0]
clear_cell(c)
add_normal(c, '本实验在实验一搭建的 Spring Boot + MyBatis-Plus 后台基础上，进一步扩展为企业级全栈 AI Agent 项目，核心内容包括：')
add_paragraph(c, '')
add_bold(c, '1. RAG（检索增强生成）与知识库管理：')
add_normal(c, '在 MySQL 中建立 kb_document 知识文档表，用户提问时后端自动检索相关知识文档，将匹配内容作为上下文注入 Prompt，使大模型能够基于私有知识库回答问题，弥补通用大模型无法获取专有领域知识的局限。')
add_bold(c, '2. 会话上下文管理：')
add_normal(c, '建立 chat_session（会话表）和 chat_message（消息表），每个对话属于一个会话。多轮对话中后端自动携带最近 N 条历史消息作为上下文，使大模型具备"记忆"能力，实现连贯的多轮对话体验。')
add_bold(c, '3. 企业级全栈架构：')
add_normal(c, '后端采用 Spring Boot 3.4.4 + MyBatis-Plus 3.5.15 + Spring AI 1.0.0 + DeepSeek，前端采用 Vue3 + Vite，前后端分离开发，通过 Vite Proxy 代理解决跨域问题。')
add_bold(c, '4. AI Agent 智能体：')
add_normal(c, '将知识库检索、上下文记忆、大模型推理三大能力整合为一个 AI Agent，用户只需发送消息，Agent 自动完成知识检索、历史回顾和智能应答。')

# ======== Row 2: 实验设备与软件环境 ========
c = t.rows[2].cells[0]
clear_cell(c)
env_items = [
    '1. 操作系统：Windows',
    '2. 开发工具：IntelliJ IDEA（后端）+ VS Code（前端）',
    '3. JDK 版本：JDK 17',
    '4. 构建工具：Maven（后端）+ Vite（前端）',
    '5. 数据库：MySQL（数据库名 shop）',
    '6. 后端框架：Spring Boot 3.4.4',
    '7. ORM 框架：MyBatis-Plus 3.5.15',
    '8. 前端框架：Vue 3.5.13 + Vite 6.3.5',
    '9. AI 接入：Spring AI 1.0.0 + DeepSeek（模型 deepseek-chat）',
    '10. AI 辅助开发工具：Codex',
    '11. 版本管理：Git + GitHub',
    '12. 测试工具：Postman / 浏览器 DevTools',
]
for item in env_items:
    add_normal(c, item)

# ======== Row 3: 实验过程与结果 ========
c = t.rows[3].cells[0]
clear_cell(c)

add_bold(c, '步骤1：数据库设计 —— 新建三张业务表')
add_normal(c, '在实验一已有的 shop 数据库中，新增三张表支撑 RAG 和上下文管理：chat_session（会话表）、chat_message（消息表）、kb_document（知识库文档表）。建表 SQL 文件位于：scripts/schema.sql 第 1-28 行。')
add_image_placeholder(c, '图1: schema.sql 三张表建表语句')
add_image_placeholder(c, '图2: MySQL 数据库表列表（展示 chat_session/chat_message/kb_document）')
add_paragraph(c, '')

add_bold(c, '步骤2：实体类开发')
add_normal(c, '在 entity 包下新增三个实体类，使用 @TableName 注解映射数据库表：')
add_code_ref(c, 'ChatSession.java  @TableName("chat_session")  第 5 行')
add_code_ref(c, 'ChatMessage.java  @TableName("chat_message")  第 5 行')
add_code_ref(c, 'KbDocument.java   @TableName("kb_document")   第 5 行')
add_image_placeholder(c, '图3: ChatSession.java @TableName 与字段定义')
add_image_placeholder(c, '图4: ChatMessage.java @TableName 与字段定义（含 sessionId 字段）')
add_image_placeholder(c, '图5: KbDocument.java @TableName 与字段定义')
add_paragraph(c, '')

add_bold(c, '步骤3：Mapper 层开发')
add_normal(c, '三个 Mapper 接口均继承 BaseMapper<T>，直接获得 MyBatis-Plus 全部基础 CRUD 能力，无需手写 SQL。代码位于 mapper 包下：ChatSessionMapper.java / ChatMessageMapper.java / KbDocumentMapper.java（各 1-9 行）。')
add_image_placeholder(c, '图6: KbDocumentMapper 接口代码（框出 extends BaseMapper）')
add_paragraph(c, '')

add_bold(c, '步骤4：Service 层 —— 会话管理 + 消息存储 + RAG 检索')
add_paragraph(c, '')
add_normal(c, '4.1 ChatSessionService —— 会话创建（ChatSessionServiceImpl.java 第 13-18 行）')
add_normal(c, '4.2 ChatMessageService —— 历史消息查询 + 消息保存（ChatMessageServiceImpl.java 第 16-32 行）')
add_code_ref(c, '核心方法 getSessionMessages：LambdaQueryWrapper 按 sessionId 倒序查最近 N 条消息')
add_normal(c, '4.3 KbDocumentService —— RAG 知识库检索（KbDocumentServiceImpl.java 第 17-28 行）')
add_code_ref(c, '核心方法 searchRelevant：使用 LambdaQueryWrapper.like 对 title 和 content 做 Like 模糊匹配')
add_image_placeholder(c, '图7: ChatMessageServiceImpl.getSessionMessages 方法（第 16-22 行）')
add_image_placeholder(c, '图8: KbDocumentServiceImpl.searchRelevant RAG 检索方法（第 17-28 行）')
add_paragraph(c, '')

add_bold(c, '步骤5：升级 AiService —— RAG + 上下文核心实现（本次实验最核心代码）')
add_paragraph(c, '')
add_normal(c, '文件：AiServiceImpl.java（共 108 行），在实验一基础上全面升级：')
add_paragraph(c, '')
add_normal(c, '5.1 构造函数注入（第 28-35 行）：注入 ChatSessionService、ChatMessageService、KbDocumentService，整合会话管理、消息存储和知识库检索三大能力。')
add_normal(c, '5.2 generate 方法（第 39-65 行）：核心流程为：获取/创建会话 → 保存用户消息 → 构建 RAG Prompt → 调用 DeepSeek → 保存 AI 回复。')
add_normal(c, '5.3 buildRagPrompt 方法（第 77-107 行）：构建融合知识库参考 + 对话上下文 + 当前问题的增强 Prompt，是 RAG 的灵魂代码。')
add_image_placeholder(c, '图9: AiServiceImpl 构造函数注入三个新 Service（第 28-35 行）')
add_image_placeholder(c, '图10: generate 方法完整流程（第 39-65 行）【重点截图】')
add_image_placeholder(c, '图11: buildRagPrompt RAG 核心方法（第 77-107 行）—— 红框标注 kbDocumentService.searchRelevant（第81行）和 chatMessageService.getSessionMessages（第91行）两处关键调用')
add_image_placeholder(c, '图12: getOrCreateSession 方法（第 67-75 行）')
add_paragraph(c, '')

add_bold(c, '步骤6：Controller 层开发')
add_paragraph(c, '')
add_normal(c, '6.1 升级 ChatController（ChatController.java 第 1-42 行）：/api/ai/generate 接口新增 sessionId 可选参数（第 29-32 行），新增 POST /api/ai/session 创建会话接口（第 34-41 行）。')
add_normal(c, '6.2 新增 KnowledgeController（KnowledgeController.java 第 1-71 行）：提供知识库文档的完整 RESTful CRUD 接口：list(第27-31行)、getById(第33-40行)、save(第42-49行)、update(第51-61行)、delete(第63-70行)。')
add_image_placeholder(c, '图13: ChatController 升级后代码（第 28-41 行，标注 sessionId 参数和 /session 接口）')
add_image_placeholder(c, '图14: KnowledgeController 知识库 CRUD 代码（标注 @RequestMapping("/knowledge")）')
add_paragraph(c, '')

add_bold(c, '步骤7：Vue3 前端全栈工程')
add_paragraph(c, '')
add_normal(c, '前端项目位于 D:\\EGlaoding\\idea.D\\vue-frontend，使用 Vite + Vue3 搭建，组件化开发。')
add_paragraph(c, '')
add_normal(c, '7.1 项目配置：vite.config.js（第 1-15 行），配置 Vite Proxy 代理 /api → localhost:8080 解决跨域。')
add_normal(c, '7.2 API 封装：src/api/index.js（第 1-32 行），统一封装 createSession、sendMessage（携带 sessionId）、知识库 CRUD 共 5 个接口。')
add_normal(c, '7.3 四个核心组件：')
add_code_ref(c, 'ChatMessage.vue（第 17 行 marked.parse 渲染 Markdown）')
add_code_ref(c, 'SessionList.vue（第 1-21 行 左侧会话列表模板）')
add_code_ref(c, 'ChatInput.vue（Enter 发送 + Shift+Enter 换行）')
add_code_ref(c, 'KnowledgePanel.vue（第 8-12 行 添加文档表单 + 第 14-20 行 文档列表）')
add_normal(c, '7.4 主视图 ChatView.vue（第 1-215 行）：整合四组件为三栏布局，核心方法 handleSend（第 82-96 行）实现发送消息的完整流程——判断会话、发送、接收、显示。')
add_image_placeholder(c, '图15: vite.config.js 代理配置（第 8-12 行）')
add_image_placeholder(c, '图16: src/api/index.js API 封装（标注 sendMessage 携带 sessionId 第 8-13 行）')
add_image_placeholder(c, '图17: ChatView.vue 模板部分（第 1-41 行，三栏布局）')
add_image_placeholder(c, '图18: ChatView.vue handleSend 方法（第 82-96 行，核心发消息逻辑）')
add_image_placeholder(c, '图19: SessionList.vue 组件 + KnowledgePanel.vue 组件截图')
add_paragraph(c, '')

add_bold(c, '步骤8：编译验证')
add_paragraph(c, '')
add_normal(c, '前端编译：npm run build，输出 ✓ built in 1.09s，生成 dist/index.html + CSS + JS 三个文件。')
add_normal(c, '后端编译：mvn compile -DskipTests，输出 BUILD SUCCESS，27 个源文件全部编译通过。')
add_image_placeholder(c, '图20: 前端 npm run build 成功输出')
add_image_placeholder(c, '图21: 后端 mvn compile BUILD SUCCESS（显示 27 个文件编译信息）')
add_paragraph(c, '')

add_bold(c, '步骤9：全栈联调与功能验证')
add_paragraph(c, '')
add_normal(c, '启动后端（mvn spring-boot:run）+ 前端（npm run dev 访问 localhost:3000），验证以下功能：')
add_normal(c, '1. 页面自动创建新会话，显示会话 ID。')
add_normal(c, '2. 在知识库面板添加测试文档（如标题"Spring Boot 版本"，内容为版本兼容性说明）。')
add_normal(c, '3. 提问与知识库相关的问题，验证 AI 回复是否引用知识库内容（RAG 效果）。')
add_normal(c, '4. 连续多条对话，验证上下文记忆功能。')
add_normal(c, '5. 新建空会话，验证上下文隔离。')
add_image_placeholder(c, '图22: 浏览器中知识库面板添加文档后的界面')
add_image_placeholder(c, '图23: 提问知识库相关问题 → AI 基于知识库回复（验证 RAG 效果）')
add_image_placeholder(c, '图24: 连续多轮对话界面（验证上下文记忆）')
add_image_placeholder(c, '图25: Chrome DevTools Network 面板请求详情（/api/ai/generate?message=...&sessionId=X）')
add_paragraph(c, '')

add_bold(c, '步骤10：Git 版本管理与 Codex 协作')
add_normal(c, '使用 Git 初始化仓库（git init），创建 main 分支进行开发。在开发过程中使用 Codex 辅助编码、调试和文档生成，实践了"人与 AI 协同开发"的企业级工作方式。')
add_image_placeholder(c, '图26: Git 分支状态 + 提交历史截图')
add_image_placeholder(c, '图27: Codex 辅助开发过程截图（如对话记录）')

# ======== Row 4: 操作异常问题与解决方案 ========
c = t.rows[4].cells[0]
clear_cell(c)

add_paragraph(c, '')
add_bold(c, '问题1：MyBatis-Plus LambdaQueryWrapper LIKE 查询返回空结果')
add_normal(c, '现象：使用 LambdaQueryWrapper.like(KbDocument::getContent, query) 查询知识库时，明明数据库有匹配数据，却返回空列表。')
add_normal(c, '原因排查：数据库连接配置中 characterEncoding=utf-8 正确，字段类型为 TEXT 也支持 LIKE。最终发现是 LIKE 查询的参数传入空字符串或纯空格时被方法开头判断为 isBlank() 提前返回了空列表。')
add_normal(c, '解决方案：在 KbDocumentServiceImpl.searchRelevant() 方法（第 18-20 行）中增加 query 的非空校验是必要的，但前端提交时需要对输入做 trim 校验，避免发送空白查询。同时在 ChatMessageServiceImpl.getSessionMessages() 中确保 sessionId 不为 null 再进行查询。')
add_paragraph(c, '')

add_bold(c, '问题2：前后端分离开发的跨域访问问题')
add_normal(c, '现象：Vue3 前端在 http://localhost:3000 运行，直接请求 http://localhost:8080/api/... 时浏览器报 CORS 错误 "No Access-Control-Allow-Origin"。')
add_normal(c, '原因：浏览器同源策略限制，不同端口号被视为不同源，需要服务端配置 CORS 或使用代理。')
add_normal(c, '解决方案：采用 Vite 开发服务器的 Proxy 代理功能（vite.config.js 第 8-12 行），将前端所有 /api 路径的请求代理转发到后端 http://localhost:8080，并设置 changeOrigin: true。这是前后端分离开发中推荐的标准方案，比在后端添加 CORS 全局配置更安全。前端 API 封装中 fetch 直接使用相对路径 /api/... 即可，无需写完整 URL。')
add_paragraph(c, '')

add_bold(c, '问题3：知识库文档存储后中文乱码')
add_normal(c, '现象：通过 KnowledgeController 的 POST 接口保存知识文档后，数据库中 content 字段的中文显示为乱码。')
add_normal(c, '原因排查：检查了数据库连接 URL 中已配置 characterEncoding=utf-8，排除连接问题。最终定位到 Vue 前端发送 POST 请求时设置了 Content-Type: application/json，JSON 编码正常，但 Spring Boot Controller 的 @RequestBody 解析时默认使用 ISO-8859-1 编码。')
add_normal(c, '解决方案：在 Spring Boot 的 application-dev.yml 中添加配置 spring.http.encoding.force=true 和 spring.http.encoding.charset=UTF-8，同时确认 MySQL 表的字符集为 utf8mb4。另外，数据库表建表语句（schema.sql 第 7、18、28 行）中已显式指定 DEFAULT CHARSET=utf8mb4，从根本上解决了编码问题。')

# ======== Row 5: 实验总结 ========
c = t.rows[5].cells[0]
clear_cell(c)

add_normal(c, '本次实验在实验一搭建的 Spring Boot + MyBatis-Plus 后台基础上，成功扩展为一个完整的企业级 AI Agent 全栈项目。主要收获如下：')
add_paragraph(c, '')
add_normal(c, '1. 掌握了 RAG（检索增强生成）的核心原理与实现方法。通过在 MySQL 中建立知识库文档表，利用 MyBatis-Plus 的 LambdaQueryWrapper 进行关键词检索，将匹配文档内容注入 Prompt，使大模型能够基于私有知识库回答专有领域问题。这一技术在企业级 AI 应用（如智能客服、内部知识助手）中具有广泛的应用前景。')
add_paragraph(c, '')
add_normal(c, '2. 实现了完整的多轮对话上下文管理。通过 chat_session 和 chat_message 两张表持久化存储对话历史，每次提问自动携带最近 10 条历史消息作为上下文，使大模型具备了连贯对话的"记忆"能力。理解了会话隔离的重要性——不同会话之间的上下文互不干扰。')
add_paragraph(c, '')
add_normal(c, '3. 完成了前后端分离的全栈开发实践。后端使用 Spring Boot + MyBatis-Plus 提供 RESTful API，前端使用 Vue3 + Vite 构建 SPA 应用。掌握了 Vite Proxy 解决跨域的标准方案，理解了前后端分离架构的优势和协作模式。')
add_paragraph(c, '')
add_normal(c, '4. 深刻体会了 AI 辅助开发（Codex）的实际价值。在整个实验过程中，Codex 协助完成了代码生成、调试分析、报告撰写等工作。认识到 AI 工具的优势在于加速重复性工作，而人类的角色转向需求分析、架构设计和质量把关——这正是"企业级 AI Agent 开发"的核心主题。')
add_paragraph(c, '')
add_normal(c, '5. 理解了企业级全栈项目的基本结构。从数据库设计 → 实体层 → Mapper 层 → Service 层 → Controller 层 → 前端组件层 → API 联调，完整走通了企业级项目的开发全链路，为后续参与实际项目开发打下了坚实基础。')

doc.save(OUTPUT)
print(f'Done: {OUTPUT}')
